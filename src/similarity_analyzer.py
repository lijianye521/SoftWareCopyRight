import os
import re
import jieba
import hashlib
import numpy as np
import difflib
from docx import Document
from sklearn.feature_extraction.text import TfidfVectorizer
from sklearn.metrics.pairwise import cosine_similarity
from simhash import Simhash

class SimilarityAnalyzer:
    def __init__(self):
        """初始化相似度分析器"""
        self.vectorizer = None
    
    def extract_text_from_docx(self, file_path):
        """从Word文档中提取文本内容，包括表格和其他元素"""
        try:
            doc = Document(file_path)
            full_text = []
            
            # 提取正文段落
            for para in doc.paragraphs:
                if para.text and para.text.strip():
                    full_text.append(para.text.strip())
            
            # 提取表格内容
            for table in doc.tables:
                for row in table.rows:
                    row_text = []
                    for cell in row.cells:
                        if cell.text and cell.text.strip():
                            row_text.append(cell.text.strip())
                    if row_text:
                        full_text.append(" ".join(row_text))
            
            # 合并所有文本
            return "\n".join(full_text)
        except Exception as e:
            print(f"无法读取文件 {file_path}: {e}")
            return ""
    
    def clean_text(self, text):
        """清理文本，保留中文标点和重要符号"""
        # 移除多余空格
        text = re.sub(r'\s+', ' ', text)
        # 移除无意义的特殊字符，但保留中文标点
        text = re.sub(r'[^\w\s\u4e00-\u9fff，。、；：""（）《》？！]', '', text)
        return text.strip()
    
    def segment_chinese_text(self, text):
        """使用结巴分词对中文文本进行分词"""
        # 分词
        words = jieba.cut(text)
        return " ".join(words)
    
    def calculate_md5(self, text):
        """计算文本的MD5哈希值，用于精确比较"""
        return hashlib.md5(text.encode('utf-8')).hexdigest()
    
    def calculate_simhash(self, text):
        """计算文本的SimHash值，用于近似比较"""
        return Simhash(text.split()).value
    
    def analyze_similarity(self, file_paths):
        """分析多个文件之间的相似度，使用多种方法"""
        if len(file_paths) < 2:
            return [], []
        
        # 提取文件名（不带路径和扩展名）
        file_names = [os.path.splitext(os.path.basename(path))[0] for path in file_paths]
        
        # 从文件中提取文本
        original_texts = []
        cleaned_texts = []
        segmented_texts = []
        md5_hashes = []
        simhashes = []
        
        for path in file_paths:
            # 提取原始文本
            original_text = self.extract_text_from_docx(path)
            original_texts.append(original_text)
            
            # 清理文本
            cleaned_text = self.clean_text(original_text)
            cleaned_texts.append(cleaned_text)
            
            # 分词
            segmented_text = self.segment_chinese_text(cleaned_text)
            segmented_texts.append(segmented_text)
            
            # 计算MD5哈希
            md5_hashes.append(self.calculate_md5(cleaned_text))
            
            # 计算SimHash
            simhashes.append(self.calculate_simhash(segmented_text))
        
        # 使用TF-IDF向量化文本
        self.vectorizer = TfidfVectorizer(
            analyzer='word',
            token_pattern=r'(?u)\b\w+\b',
            min_df=0.0,  # 降低min_df值，允许所有词项
            max_df=1.0,  # 提高max_df值，允许所有词项
            ngram_range=(1, 3)  # 使用1-gram、2-gram和3-gram
        )
        
        try:
            tfidf_matrix = self.vectorizer.fit_transform(segmented_texts)
            # 计算余弦相似度
            cosine_sim_matrix = cosine_similarity(tfidf_matrix)
        except Exception as e:
            print(f"TF-IDF向量化失败: {e}")
            # 创建一个全零的相似度矩阵作为备选
            cosine_sim_matrix = np.zeros((len(file_paths), len(file_paths)))
        
        # 格式化结果
        results = []
        for i in range(len(file_paths)):
            for j in range(i+1, len(file_paths)):
                # 计算各种相似度指标
                cosine_sim = cosine_sim_matrix[i][j] * 100  # 转换为百分比
                
                # 计算序列相似度 (difflib)
                seq_sim = difflib.SequenceMatcher(None, cleaned_texts[i], cleaned_texts[j]).ratio() * 100
                
                # 检查MD5哈希是否完全相同
                exact_match = md5_hashes[i] == md5_hashes[j]
                
                # 计算SimHash距离 (数值越小表示越相似)
                simhash_distance = bin(simhashes[i] ^ simhashes[j]).count('1')
                simhash_sim = 100 - (simhash_distance / 64 * 100)  # 转换为相似度百分比
                
                # 综合相似度 (加权平均)
                combined_sim = 0.4 * cosine_sim + 0.3 * seq_sim + 0.3 * simhash_sim
                
                # 如果MD5完全匹配，则相似度为100%
                if exact_match:
                    combined_sim = 100.0
                
                results.append({
                    'file1': file_names[i],
                    'file2': file_names[j],
                    'similarity': combined_sim,
                    'cosine_similarity': cosine_sim,
                    'sequence_similarity': seq_sim,
                    'simhash_similarity': simhash_sim,
                    'exact_match': exact_match
                })
        
        # 按相似度降序排序
        results.sort(key=lambda x: x['similarity'], reverse=True)
        
        return results, cosine_sim_matrix
    
    def get_similarity_report(self, file_paths):
        """生成相似度报告"""
        if len(file_paths) < 2:
            return "需要至少两个文件才能进行相似度分析。"
        
        # 分析相似度
        results, _ = self.analyze_similarity(file_paths)
        
        # 生成报告
        report = "文件相似度分析报告:\n\n"
        
        # 添加详细的相似度结果
        for result in results:
            report += f"{result['file1']} 与 {result['file2']} 的相似度: {result['similarity']:.2f}%\n"
            
            if result['exact_match']:
                report += "  【警告】这两个文件内容完全相同！\n"
                
            report += f"  - 余弦相似度: {result['cosine_similarity']:.2f}%\n"
            report += f"  - 序列相似度: {result['sequence_similarity']:.2f}%\n"
            report += f"  - SimHash相似度: {result['simhash_similarity']:.2f}%\n\n"
        
        # 找出最相似的文件对
        if results:
            most_similar = results[0]
            report += f"\n最相似的文件对: {most_similar['file1']} 和 {most_similar['file2']} ({most_similar['similarity']:.2f}%)"
            if most_similar['exact_match']:
                report += " - 内容完全相同！"
        
        return report 