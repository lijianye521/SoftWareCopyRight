import tkinter as tk
from tkinter import ttk, filedialog, messagebox, font, Toplevel, scrolledtext
from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import qn as shared_qn
import os
import pathlib
import chardet
import pathspec
import json
import configparser
from PIL import Image, ImageTk, ImageDraw  # 用于美化界面图标
import sv_ttk  # 用于现代化Tkinter主题

class CodeToWordApp:
    # 添加创建圆角图像的方法
    def create_rounded_button_images(self):
        """创建圆角按钮图像"""
        # 这个方法会在初始化时被调用，用于创建圆角按钮图像
        self.rounded_images = {}
        
    def create_rounded_rectangle(self, width, height, radius, fill_color):
        """创建圆角矩形图像"""
        image = Image.new("RGBA", (width, height), (0, 0, 0, 0))
        draw = ImageDraw.Draw(image)
        draw.rounded_rectangle([(0, 0), (width, height)], radius, fill=fill_color)
        return ImageTk.PhotoImage(image)
        
    def __init__(self, root):
        self.root = root
        self.root.title("软件著作权源代码文档生成器")
        self.root.geometry("900x950") # 增大窗口宽度以适应新设计
        
        # 应用Sun Valley主题 - 一个现代化的Tkinter主题，提供圆角和阴影效果
        sv_ttk.set_theme("light")
        
        # 设置应用主题和样式
        self.style = ttk.Style()
        
        # 定义统一的楷体字体和大小
        label_font = ('楷体', 14)
        entry_font = ('楷体', 12)  # 输入框字体改为12号
        button_font = ('楷体', 14)
        
        # 自定义样式
        self.style.configure('TLabel', font=label_font)
        self.style.configure('TEntry', font=entry_font)
        self.style.configure('TButton', font=button_font)
        self.style.configure('TLabelframe.Label', font=label_font)
        self.style.configure('TCheckbutton', font=label_font)
        self.style.configure('TCombobox', font=entry_font)
        
        self.style.configure('Small.TButton', font=button_font)
        self.style.configure('Add.TButton', font=button_font, padding=(10, 5))
        self.style.configure('Generate.TButton', font=('楷体', 16, 'bold'), padding=(20, 10))
        self.style.configure('Title.TLabel', font=('楷体', 24, 'bold'), foreground="#4361ee")
        self.style.configure('Subtitle.TLabel', font=('楷体', 16))
        self.style.configure('Info.TLabel', font=('楷体', 14))
        
        # 创建圆角按钮图像
        self.create_rounded_button_images()
        
        # 设置背景颜色 - 使用更柔和的颜色
        self.root.configure(bg="#f8f9fa")
        
        # 预定义文件类型模式
        self.file_type_modes = {
            "前端模式": ["json", "js", "jsx", "ts", "tsx", "html", "css", "scss", "vue"],
            "后端模式": ["java", "xml", "py", "php", "cs", "go", "rb"],
            "混合模式": ["json", "js", "jsx", "ts", "tsx", "html", "css", "scss", "vue", "java", "xml", "py", "php"]
        }
        
        # 加载用户自定义模式
        self.config = configparser.ConfigParser()
        self.config_file = os.path.join(os.path.dirname(os.path.abspath(__file__)), "config.ini")
        self.load_config()
        
        # 项目路径列表
        self.project_paths = []
        
        # 创建主框架，使用滚动条
        outer_frame = ttk.Frame(root)
        outer_frame.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S))
        outer_frame.columnconfigure(0, weight=1)
        outer_frame.rowconfigure(0, weight=1)
        
        # 创建Canvas和滚动条
        canvas = tk.Canvas(outer_frame, bg="#f8f9fa", highlightthickness=0)
        scrollbar = ttk.Scrollbar(outer_frame, orient="vertical", command=canvas.yview)
        canvas.configure(yscrollcommand=scrollbar.set)
        
        scrollbar.grid(row=0, column=1, sticky=(tk.N, tk.S))
        canvas.grid(row=0, column=0, sticky=(tk.W, tk.E, tk.N, tk.S), padx=10, pady=10)
        
        # 主内容框架
        main_frame = ttk.Frame(canvas, padding="30")
        main_frame.columnconfigure(1, weight=1)
        
        # 创建窗口用于放置main_frame
        canvas_window = canvas.create_window((0, 0), window=main_frame, anchor="nw")
        
        # 配置Canvas滚动区域
        def configure_scroll_region(event):
            canvas.configure(scrollregion=canvas.bbox("all"))
            # 设置canvas宽度与main_frame相同
            canvas.itemconfigure(canvas_window, width=event.width)
        
        main_frame.bind("<Configure>", configure_scroll_region)
        canvas.bind("<Configure>", lambda e: canvas.itemconfigure(canvas_window, width=e.width))
        
        # 绑定鼠标滚轮事件到Canvas
        def _on_mousewheel(event):
            canvas.yview_scroll(int(-1*(event.delta/120)), "units")
        
        canvas.bind_all("<MouseWheel>", _on_mousewheel)
        

        
        # 标题和图标
        header_frame = ttk.Frame(main_frame)
        header_frame.grid(row=0, column=0, columnspan=3, pady=(0, 30))
        
        # 创建标题
        title_label = ttk.Label(
            header_frame, 
            text="软件源代码导出工具", 
            style='Title.TLabel'
        )
        title_label.pack(pady=15)
        
        # 副标题
        subtitle_label = ttk.Label(
            header_frame, 
            text="用于软件著作权申请的源代码文档生成器", 
            style='Subtitle.TLabel'
        )
        subtitle_label.pack(pady=5)
        
        # 项目路径选择框架 - 可动态添加多个路径
        paths_frame = ttk.LabelFrame(main_frame, text="项目路径", padding=15)
        paths_frame.grid(row=1, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=15, padx=5)
        paths_frame.columnconfigure(0, weight=1)
        
        # 初始化第一个路径选择行
        self.path_vars = []
        self.path_frames = []
        self.add_path_row()
        
        # 添加路径按钮
        add_path_btn = ttk.Button(
            paths_frame, 
            text="+ 添加路径", 
            command=self.add_path_row,
            style='Add.TButton'
        )
        add_path_btn.grid(row=100, column=0, sticky=tk.W, pady=15)
        
        # Gitignore 文件选择框架
        gitignore_frame = ttk.LabelFrame(main_frame, text="Gitignore 设置", padding=15)
        gitignore_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=15, padx=5)
        gitignore_frame.columnconfigure(1, weight=1)
        
        ttk.Label(gitignore_frame, text="Gitignore 文件:").grid(row=0, column=0, sticky=tk.W, pady=5)
        self.gitignore_path_var = tk.StringVar()
        gitignore_entry = ttk.Entry(gitignore_frame, textvariable=self.gitignore_path_var, width=50)
        gitignore_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=5, padx=5)
        gitignore_browse_btn = ttk.Button(gitignore_frame, text="浏览...", command=self.choose_gitignore_file)
        gitignore_browse_btn.grid(row=0, column=2, padx=10, pady=5)
        
        # 文件类型模式选择框架
        file_types_frame = ttk.LabelFrame(main_frame, text="文件类型设置", padding=15)
        file_types_frame.grid(row=3, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=15, padx=5)
        file_types_frame.columnconfigure(1, weight=1)
        
        # 获取所有模式名称
        mode_names = list(self.file_type_modes.keys()) + ["其他模式"]
        self.file_mode_var = tk.StringVar(value=mode_names[0])
        
        # 模式选择标签和下拉框
        ttk.Label(file_types_frame, text="选择模式:").grid(row=0, column=0, sticky=tk.W, pady=10)
        
        # 模式选择下拉框
        mode_selector = ttk.Combobox(
            file_types_frame, 
            textvariable=self.file_mode_var, 
            values=mode_names, 
            state="readonly", 
            width=20
        )
        mode_selector.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=10, padx=5)
        mode_selector.bind("<<ComboboxSelected>>", self.on_mode_selected)
        
        # 文件类型标签和输入框
        ttk.Label(file_types_frame, text="文件类型:").grid(row=1, column=0, sticky=tk.W, pady=10)
        
        # 自定义文件类型输入框
        self.extension_var = tk.StringVar()
        self.extension_entry = ttk.Entry(file_types_frame, textvariable=self.extension_var)
        self.extension_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=10, padx=5)
        
        # 添加自定义模式按钮
        save_mode_btn = ttk.Button(
            file_types_frame, 
            text="保存为自定义模式", 
            command=self.save_custom_mode
        )
        save_mode_btn.grid(row=2, column=1, sticky=tk.E, pady=(10, 0), padx=5)
        
        # 初始化文件类型
        self.update_extension_field()
        
        # 文档设置框架
        doc_settings_frame = ttk.LabelFrame(main_frame, text="文档设置", padding=15)
        doc_settings_frame.grid(row=4, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=15, padx=5)
        doc_settings_frame.columnconfigure(1, weight=1)
        
        # 字体选择
        ttk.Label(doc_settings_frame, text="字体:").grid(row=0, column=0, sticky=tk.W, pady=10)
        self.font_var = tk.StringVar(value="微软雅黑")
        
        font_display = ttk.Entry(doc_settings_frame, textvariable=self.font_var, state='readonly')
        font_display.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=10, padx=5)
        
        font_btn = ttk.Button(doc_settings_frame, text="选择字体", command=lambda: self.choose_font(self.font_var))
        font_btn.grid(row=0, column=2, padx=5, pady=10)
        
        # 每页行数设置
        ttk.Label(doc_settings_frame, text="每页行数:").grid(row=1, column=0, sticky=tk.W, pady=10)
        self.lines_per_page_var = tk.StringVar(value="55")
        lines_per_page_entry = ttk.Entry(doc_settings_frame, textvariable=self.lines_per_page_var)
        lines_per_page_entry.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=10, padx=5)
        
        # 应用名称和版本号设置
        ttk.Label(doc_settings_frame, text="应用名称:").grid(row=2, column=0, sticky=tk.W, pady=10)
        self.app_name_var = tk.StringVar(value="AI简报")
        self.app_name_font_var = tk.StringVar(value="微软雅黑")
        
        app_name_entry = ttk.Entry(doc_settings_frame, textvariable=self.app_name_var)
        app_name_entry.grid(row=2, column=1, sticky=(tk.W, tk.E), pady=10, padx=5)
        
        app_font_btn = ttk.Button(doc_settings_frame, text="字体", command=lambda: self.choose_font(self.app_name_font_var), width=5)
        app_font_btn.grid(row=2, column=2, padx=5, pady=10)
        
        # 版本号和字体选择按钮
        ttk.Label(doc_settings_frame, text="版本号:").grid(row=3, column=0, sticky=tk.W, pady=10)
        self.version_var = tk.StringVar(value="V1.0")
        self.version_font_var = tk.StringVar(value="微软雅黑")
        
        version_entry = ttk.Entry(doc_settings_frame, textvariable=self.version_var)
        version_entry.grid(row=3, column=1, sticky=(tk.W, tk.E), pady=10, padx=5)
        
        version_font_btn = ttk.Button(doc_settings_frame, text="字体", command=lambda: self.choose_font(self.version_font_var), width=5)
        version_font_btn.grid(row=3, column=2, padx=5, pady=10)
        
        # 输出设置框架
        output_frame = ttk.LabelFrame(main_frame, text="输出设置", padding=15)
        output_frame.grid(row=5, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=15, padx=5)
        output_frame.columnconfigure(1, weight=1)
        
        # 输出文档路径
        ttk.Label(output_frame, text="输出路径:").grid(row=0, column=0, sticky=tk.W, pady=10)
        self.output_path_var = tk.StringVar()
        
        output_path_entry = ttk.Entry(output_frame, textvariable=self.output_path_var)
        output_path_entry.grid(row=0, column=1, sticky=(tk.W, tk.E), pady=10, padx=5)
        
        output_path_btn = ttk.Button(output_frame, text="浏览...", command=self.choose_output_directory)
        output_path_btn.grid(row=0, column=2, padx=5, pady=10)
        
        # 生成模式
        ttk.Label(output_frame, text="生成模式:").grid(row=1, column=0, sticky=tk.W, pady=10)
        self.mode_var = tk.StringVar(value="默认模式")
        mode_selector = ttk.Combobox(output_frame, textvariable=self.mode_var, values=["默认模式", "标准模式"], state="readonly")
        mode_selector.grid(row=1, column=1, sticky=(tk.W, tk.E), pady=10, padx=5)
        
        # 选项框架
        options_frame = ttk.LabelFrame(output_frame, text="文档选项", padding=10)
        options_frame.grid(row=2, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=(15, 0))
        options_frame.columnconfigure(0, weight=1)
        options_frame.columnconfigure(1, weight=1)
        
        # 生成目录选项
        self.generate_toc_var = tk.BooleanVar(value=False)
        toc_check = ttk.Checkbutton(options_frame, text="生成目录", variable=self.generate_toc_var)
        toc_check.grid(row=0, column=0, sticky=tk.W, padx=20, pady=10)
        
        # 显示行号选项
        self.show_line_numbers_var = tk.BooleanVar(value=False)
        line_numbers_check = ttk.Checkbutton(options_frame, text="显示行号", variable=self.show_line_numbers_var)
        line_numbers_check.grid(row=0, column=1, sticky=tk.W, padx=20, pady=10)
        
        # 模式说明
        info_frame = ttk.LabelFrame(main_frame, text="模式说明", padding=15)
        info_frame.grid(row=6, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=15, padx=5)
        
        standard_mode_info = ttk.Label(
            info_frame, 
            text="• 默认模式：生成所有源代码文件。\n• 标准模式：用于软件著作权申请，将生成源代码文件的前30页和后30页。\n  要求总代码量超过60页，否则将按默认模式生成。", 
            style='Info.TLabel',
            wraplength=750, 
            justify=tk.LEFT
        )
        standard_mode_info.pack(anchor=tk.W, pady=10)
        
        # 生成按钮
        button_frame = ttk.Frame(main_frame)
        button_frame.grid(row=7, column=0, columnspan=3, pady=30)
        
        # 创建一个带有阴影效果的生成按钮
        generate_btn = ttk.Button(
            button_frame, 
            text="生成文档", 
            command=self.generate_document,
            style='Generate.TButton'
        )
        generate_btn.pack(padx=10)
        
        # 底部信息栏
        footer_frame = ttk.Frame(main_frame)
        footer_frame.grid(row=8, column=0, columnspan=3, sticky=(tk.W, tk.E), pady=10)
        footer_frame.columnconfigure(1, weight=1)
        
        # 版本信息
        version_label = ttk.Label(footer_frame, text="v1.2.0")
        version_label.grid(row=0, column=0, sticky=tk.W)
        
        # 作者信息
        author_label = ttk.Label(footer_frame, text="作者: lijianye")
        author_label.grid(row=0, column=1, sticky=tk.E)
        
        # 状态变量（不显示状态框）
        self.status_var = tk.StringVar()
        self.progress_window = None
        
        # 使窗口可调整大小
        self.root.columnconfigure(0, weight=1)
        self.root.rowconfigure(0, weight=1)
        
    def add_path_row(self):
        """添加一个新的项目路径选择行"""
        # 创建新的路径变量
        path_var = tk.StringVar()
        self.path_vars.append(path_var)
        
        # 获取路径框架
        if not hasattr(self, 'paths_frame'):
            # 第一次调用时，保存路径框架的引用
            for widget in self.root.winfo_children():
                if isinstance(widget, ttk.Frame):  # 外层框架
                    for child in widget.winfo_children():
                        if isinstance(child, tk.Canvas):  # Canvas
                            frame_id = child.find_withtag("all")[0]
                            main_frame = child.itemcget(frame_id, "window")
                            if main_frame:
                                main_frame = child.nametowidget(main_frame)
                                for mf_child in main_frame.winfo_children():
                                    if isinstance(mf_child, ttk.Labelframe) and mf_child.cget("text") == "项目路径":
                                        self.paths_frame = mf_child
                                        break
        
        row_idx = len(self.path_frames)
        path_frame = ttk.Frame(self.paths_frame)
        path_frame.grid(row=row_idx, column=0, sticky=(tk.W, tk.E), pady=5)
        path_frame.columnconfigure(0, weight=1)
        
        self.path_frames.append(path_frame)
        
        # 路径输入框
        path_entry = ttk.Entry(path_frame, textvariable=path_var, width=50, font=('微软雅黑', 12))
        path_entry.grid(row=0, column=0, sticky=(tk.W, tk.E))
        
        # 浏览按钮
        browse_btn = ttk.Button(
            path_frame, 
            text="浏览...", 
            command=lambda v=path_var: self.choose_directory(v),
            style='TButton'
        )
        browse_btn.grid(row=0, column=1, padx=5)
        
        # 删除按钮（第一行不显示）
        if row_idx > 0:
            delete_btn = ttk.Button(
                path_frame,
                text="删除",
                command=lambda idx=row_idx: self.delete_path_row(idx),
                style='Small.TButton'
            )
            delete_btn.grid(row=0, column=2)
    
    def delete_path_row(self, idx):
        """删除指定索引的路径行"""
        if idx < len(self.path_frames) and idx > 0:  # 不允许删除第一行
            # 销毁UI元素
            self.path_frames[idx].destroy()
            # 从列表中移除
            self.path_frames.pop(idx)
            self.path_vars.pop(idx)
            
            # 重新排列剩余行
            for i in range(idx, len(self.path_frames)):
                self.path_frames[i].grid(row=i, column=0)
                # 更新删除按钮的命令
                delete_btn = self.path_frames[i].winfo_children()[2]
                delete_btn.configure(command=lambda idx=i: self.delete_path_row(idx))
    
    def choose_directory(self, path_var):
        """选择目录并设置到指定的StringVar中"""
        directory = filedialog.askdirectory()
        if directory:
            path_var.set(directory)
    
    def choose_output_directory(self):
        directory = filedialog.askdirectory()
        if directory:
            self.output_path_var.set(directory)
    
    def choose_gitignore_file(self):
        """打开文件对话框选择.gitignore文件"""
        filepath = filedialog.askopenfilename(
            title="选择 .gitignore 文件",
            filetypes=[("Gitignore files", ".gitignore"), ("All files", "*.*")]
        )
        if filepath:
            self.gitignore_path_var.set(filepath)
            
    def load_config(self):
        """加载配置文件"""
        try:
            if os.path.exists(self.config_file):
                self.config.read(self.config_file, encoding='utf-8')
                
                # 加载自定义模式
                if 'CustomModes' in self.config:
                    for mode_name, extensions in self.config['CustomModes'].items():
                        if mode_name not in self.file_type_modes:
                            self.file_type_modes[mode_name] = extensions.split(',')
        except Exception as e:
            print(f"加载配置文件失败: {e}")
    
    def save_config(self):
        """保存配置到文件"""
        try:
            # 确保存在CustomModes部分
            if 'CustomModes' not in self.config:
                self.config['CustomModes'] = {}
            
            # 保存自定义模式
            custom_modes = {k: ','.join(v) for k, v in self.file_type_modes.items() 
                           if k not in ["前端模式", "后端模式", "混合模式"]}
            
            self.config['CustomModes'] = custom_modes
            
            # 写入文件
            os.makedirs(os.path.dirname(self.config_file), exist_ok=True)
            with open(self.config_file, 'w', encoding='utf-8') as f:
                self.config.write(f)
        except Exception as e:
            messagebox.showerror("错误", f"保存配置失败: {e}")
    
    def on_mode_selected(self, event=None):
        """当文件类型模式被选择时更新扩展名输入框"""
        self.update_extension_field()
    
    def update_extension_field(self):
        """根据选择的模式更新扩展名输入框"""
        mode = self.file_mode_var.get()
        
        if mode == "其他模式":
            # 启用输入框
            self.extension_entry.config(state='normal')
        else:
            # 显示预设模式的扩展名
            extensions = self.file_type_modes.get(mode, [])
            self.extension_var.set(','.join(extensions))
            # 禁用输入框，除非是自定义模式
            is_custom = mode not in ["前端模式", "后端模式", "混合模式"]
            self.extension_entry.config(state='normal' if is_custom else 'readonly')
    
    def save_custom_mode(self):
        """保存当前扩展名设置为自定义模式"""
        extensions = [ext.strip() for ext in self.extension_var.get().split(',') if ext.strip()]
        if not extensions:
            messagebox.showerror("错误", "请输入至少一个文件类型")
            return
            
        # 弹出对话框让用户输入模式名称
        mode_dialog = Toplevel(self.root)
        mode_dialog.title("保存自定义模式")
        mode_dialog.geometry("400x150")
        mode_dialog.resizable(False, False)
        mode_dialog.transient(self.root)
        
        ttk.Label(mode_dialog, text="请输入模式名称:").pack(pady=(20, 10))
        
        mode_name_var = tk.StringVar()
        mode_name_entry = ttk.Entry(mode_dialog, textvariable=mode_name_var, width=30, font=('微软雅黑', 12))
        mode_name_entry.pack(padx=20, fill=tk.X)
        
        def on_save():
            mode_name = mode_name_var.get().strip()
            if not mode_name:
                messagebox.showerror("错误", "模式名称不能为空")
                return
                
            if mode_name in ["前端模式", "后端模式", "混合模式", "其他模式"]:
                messagebox.showerror("错误", "不能使用预设模式名称")
                return
                
            # 保存模式
            self.file_type_modes[mode_name] = extensions
            
            # 更新下拉框选项
            # 寻找模式选择下拉框
            for widget in self.root.winfo_children():
                if isinstance(widget, ttk.Frame):  # 外层框架
                    for child in widget.winfo_children():
                        if isinstance(child, tk.Canvas):  # Canvas
                            frame_id = child.find_withtag("all")[0]
                            main_frame = child.itemcget(frame_id, "window")
                            if main_frame:
                                main_frame = child.nametowidget(main_frame)
                                for i, mf_child in enumerate(main_frame.winfo_children()):
                                    if i == 3:  # 第4个子元素应该是文件类型模式框架
                                        mode_frame = mf_child.winfo_children()[1]  # 获取模式框架
                                        mode_selector = mode_frame.winfo_children()[0]  # 获取下拉框
                                        break
            
            # 更新下拉框
            mode_names = list(self.file_type_modes.keys()) + ["其他模式"]
            mode_selector['values'] = mode_names
            self.file_mode_var.set(mode_name)
            
            # 保存到配置文件
            self.save_config()
            
            messagebox.showinfo("成功", f"已保存模式 '{mode_name}'")
            mode_dialog.destroy()
            
        # 按钮框架
        btn_frame = ttk.Frame(mode_dialog)
        btn_frame.pack(pady=20, fill=tk.X)
        
        ttk.Button(btn_frame, text="保存", command=on_save).pack(side=tk.RIGHT, padx=10)
        ttk.Button(btn_frame, text="取消", command=mode_dialog.destroy).pack(side=tk.RIGHT, padx=10)
        
        # 聚焦输入框
        mode_name_entry.focus_set()
    
    def choose_font(self, font_var):
        """打开字体选择对话框"""
        font_dialog = Toplevel(self.root)
        font_dialog.title("选择字体")
        font_dialog.geometry("400x300")
        font_dialog.resizable(False, False)
        font_dialog.transient(self.root)
        
        # 获取系统字体列表
        font_list = list(font.families())
        font_list.sort()
        
        # 创建字体列表框
        font_listbox = tk.Listbox(font_dialog, font=('楷体', 12), height=10)
        font_listbox.pack(fill=tk.BOTH, expand=True, padx=20, pady=10)
        
        # 添加字体到列表框
        for f in font_list:
            font_listbox.insert(tk.END, f)
        
        # 如果当前字体在列表中，选中它
        current_font = font_var.get()
        if current_font in font_list:
            index = font_list.index(current_font)
            font_listbox.see(index)
            font_listbox.selection_set(index)
        
        # 确定按钮回调
        def on_select():
            selection = font_listbox.curselection()
            if selection:
                selected_font = font_list[selection[0]]
                font_var.set(selected_font)
            font_dialog.destroy()
        
        # 确定按钮
        select_btn = ttk.Button(font_dialog, text="确定", command=on_select)
        select_btn.pack(pady=10)

    def detect_encoding(self, file_path):
        """检测文件编码"""
        with open(file_path, 'rb') as f:
            result = chardet.detect(f.read())
        return result['encoding'] or 'utf-8'

    def read_file_lines(self, file_path, encoding):
        """读取文件内容，将软回车转为硬回车，移除空行，并返回行列表"""
        with open(file_path, 'r', encoding=encoding) as f:
            content = f.read()
        # 将软回车（\v 或 ^l）替换为硬回车（\n 或 ^p）
        content = content.replace('\v', '\n')
        lines = content.splitlines()
        # 只保留非空行
        return [line for line in lines if line.strip()]
    
    def apply_style_to_paragraph(self, paragraph, font_name="微软雅黑", font_size=10):
        """应用统一样式到段落"""
        # 设置段落行距为精确值
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(font_size * 1.05)  # 进一步减小行距，设为字体大小的1.05倍
        # 设置段落间距为0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            # 显式设置东亚字体
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def add_header_footer(self, doc, app_name, version, app_name_font, version_font):
        """添加页眉和页码"""
        # 获取第一节
        section = doc.sections[0]
        
        # 创建页眉
        header = section.header
        header_para = header.paragraphs[0]
        
        # 添加应用名称
        app_run = header_para.add_run(app_name)
        app_run.font.name = app_name_font
        app_run.font.size = Pt(12)
        app_run.bold = True
        # 显式设置东亚字体
        r_app = app_run._element
        r_app.rPr.rFonts.set(qn('w:eastAsia'), app_name_font)
        
        # 添加空格
        header_para.add_run(" ")
        
        # 添加版本号
        ver_run = header_para.add_run(version)
        ver_run.font.name = version_font
        ver_run.font.size = Pt(12)
        # 显式设置东亚字体
        r_ver = ver_run._element
        r_ver.rPr.rFonts.set(qn('w:eastAsia'), version_font)
        
        # 设置页眉段落格式
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # 添加页码（第n页/共m页）
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # 添加"第"
        page_text_run = footer_para.add_run("第")
        page_text_run.font.name = "微软雅黑"
        page_text_run.font.size = Pt(10)
        r_text = page_text_run._element
        r_text.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        
        # 添加当前页码域代码
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(shared_qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(shared_qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(shared_qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        # 添加"页/共"
        page_text_run2 = footer_para.add_run("页/共")
        page_text_run2.font.name = "微软雅黑"
        page_text_run2.font.size = Pt(10)
        r_text2 = page_text_run2._element
        r_text2.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        
        # 添加总页数域代码
        run = footer_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(shared_qn('w:fldCharType'), 'begin')
        run._r.append(fldChar3)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(shared_qn('xml:space'), 'preserve')
        instrText.text = ' NUMPAGES '
        run._r.append(instrText)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(shared_qn('w:fldCharType'), 'end')
        run._r.append(fldChar4)
        
        # 添加"页"
        page_text_run3 = footer_para.add_run("页")
        page_text_run3.font.name = "微软雅黑"
        page_text_run3.font.size = Pt(10)
        r_text3 = page_text_run3._element
        r_text3.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
            
    def _create_progress_window(self, total_files):
        """创建并显示进度条窗口"""
        self.progress_window = Toplevel(self.root)
        self.progress_window.title("正在生成")
        self.progress_window.geometry("400x120")
        self.progress_window.resizable(False, False)
        self.progress_window.transient(self.root) # 保持在主窗口之上
        
        ttk.Label(self.progress_window, text="正在处理文件...", font=('微软雅黑', 12)).pack(pady=10)
        
        self.progress_bar = ttk.Progressbar(self.progress_window, orient="horizontal", length=350, mode="determinate")
        self.progress_bar.pack(pady=5)
        self.progress_bar["maximum"] = total_files
        
        self.progress_label = ttk.Label(self.progress_window, text="", font=('微软雅黑', 9))
        self.progress_label.pack(pady=5)
        
        self.root.update_idletasks()

    def generate_document(self):
        # 收集所有项目路径
        paths = [var.get() for var in self.path_vars if var.get()]
        if not paths:
            messagebox.showerror("错误", "请至少选择一个项目路径")
            return
            
        # 获取文件类型
        mode = self.file_mode_var.get()
        if mode == "其他模式":
            extensions = [ext.strip() for ext in self.extension_var.get().split(',')]
        else:
            extensions = self.file_type_modes.get(mode, [])
            
        if not extensions:
            messagebox.showerror("错误", "请选择或输入文件类型")
            return
        
        font_name = self.font_var.get()
        if not font_name:
            font_name = "微软雅黑"
            
        app_name = self.app_name_var.get()
        version = self.version_var.get()
        app_name_font = self.app_name_font_var.get() or "微软雅黑"
        version_font = self.version_font_var.get() or "微软雅黑"
        
        mode = self.mode_var.get()
        
        # 获取并验证每页行数
        try:
            lines_per_page = int(self.lines_per_page_var.get())
            if lines_per_page <= 0:
                raise ValueError("每页行数必须大于0")
        except ValueError:
            messagebox.showerror("错误", "请输入有效的每页行数")
            return
            
        # 根据每页行数计算适当的字体大小（简单算法，可能需要调整）
        line_height_cm = 26.7 / lines_per_page
        font_size = int(line_height_cm * 28)
        font_size = max(8, min(12, font_size))
        
        try:
            # --- Stage 1: File Discovery & .gitignore Filtering ---
            all_files = []
            
            # 获取 .gitignore 规则
            gitignore_path_str = self.gitignore_path_var.get()
            gitignore_path = None
            spec = None
            
            if gitignore_path_str and os.path.exists(gitignore_path_str):
                gitignore_path = pathlib.Path(gitignore_path_str)
                with open(gitignore_path, 'r', encoding='utf-8') as f:
                    spec = pathspec.PathSpec.from_lines('gitwildmatch', f)
            
            # 遍历所有项目路径收集文件
            for path in paths:
                base_path = pathlib.Path(path)
                
                # 如果没有指定gitignore但项目中有，使用项目中的
                if not spec and (base_path / '.gitignore').is_file():
                    with open(base_path / '.gitignore', 'r', encoding='utf-8') as f:
                        spec = pathspec.PathSpec.from_lines('gitwildmatch', f)
                
                # 收集所有匹配的文件
                for extension in extensions:
                    extension = extension.strip('.')
                    all_files.extend(base_path.rglob(f"*.{extension}"))

            files_to_process = []
            ignored_files = []
            if spec:
                for f in all_files:
                    if not spec.match_file(str(f.relative_to(base_path))):
                        files_to_process.append(f)
                    else:
                        ignored_files.append(str(f.relative_to(base_path)))
            else:
                files_to_process = all_files

            # --- Stage 1.5: Setup Progress Bar ---
            self._create_progress_window(len(files_to_process))

            # --- Stage 2: Collect all lines from all files ---
            all_styled_lines = []
            lines_by_ext = {}

            for i, file_path in enumerate(files_to_process):
                # Update progress
                self.progress_bar['value'] = i + 1
                self.progress_label['text'] = f"正在处理: {file_path.name}"
                self.progress_window.update_idletasks()

                ext = file_path.suffix
                if ext not in lines_by_ext:
                    lines_by_ext[ext] = 0
                
                relative_path = str(file_path.relative_to(base_path))
                all_styled_lines.append((relative_path, 'path'))
                
                try:
                    encoding = self.detect_encoding(file_path)
                    lines = self.read_file_lines(file_path, encoding)
                    lines_by_ext[ext] += len(lines)
                    for line in lines:
                        all_styled_lines.append((line, 'code'))
                except Exception as e:
                    all_styled_lines.append((f"无法读取文件: {relative_path} ({e})", 'error'))

            file_count = len(files_to_process)
            
            if self.progress_window:
                self.progress_window.destroy()
                self.progress_window = None

            # --- Stage 3: Decide which lines to print based on mode ---
            lines_to_print = []
            total_lines = len(all_styled_lines)
            total_pages = total_lines / lines_per_page
            
            final_mode = mode
            if mode == "标准模式":
                if total_pages <= 60:
                    messagebox.showwarning("模式切换", f"总代码约 {total_pages:.1f} 页，不足60页，无法按标准模式生成。\n将自动切换到默认模式。")
                    final_mode = "默认模式"
                else:
                    lines_for_30_pages = int(30 * lines_per_page)
                    first_30 = all_styled_lines[:lines_for_30_pages]
                    last_30 = all_styled_lines[-lines_for_30_pages:]
                    separator_text = f"... (中间部分代码省略) ..."
                    
                    lines_to_print.extend(first_30)
                    lines_to_print.append((separator_text, 'separator'))
                    lines_to_print.extend(last_30)
            
            if final_mode == "默认模式":
                lines_to_print = all_styled_lines

            # --- Stage 4: Generate the document from the prepared line list ---
            doc = Document()
            # 设置文档页面大小和边距
            for section in doc.sections:
                section.page_height = Cm(29.7)  # A4高度
                section.page_width = Cm(21.0)   # A4宽度
                # 设置更小的页边距以容纳更多内容
                section.top_margin = Cm(0.8)    # 进一步减小上边距
                section.bottom_margin = Cm(0.8) # 进一步减小下边距
                section.left_margin = Cm(1.5)   # 减小左边距
                section.right_margin = Cm(1.5)  # 减小右边距
                # 设置页眉和页脚距离
                section.header_distance = Cm(0.3)
                section.footer_distance = Cm(0.3)
            
            # 设置文档默认字体和段落格式
            style = doc.styles['Normal']
            style.font.name = font_name
            style.font.size = Pt(font_size)
            rpr = style.element.rPr
            rpr.rFonts.set(qn('w:eastAsia'), font_name)
            style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
            style.paragraph_format.line_spacing = Pt(font_size * 1.05)
            style.paragraph_format.space_before = Pt(0)
            style.paragraph_format.space_after = Pt(0)
            
            self.add_header_footer(doc, app_name, version, app_name_font, version_font)
            
            # 如果需要生成目录
            if self.generate_toc_var.get():
                doc.add_paragraph("目录", style="Heading 1")
                doc.add_paragraph()  # 添加空行
                
                # 添加目录域代码
                paragraph = doc.add_paragraph()
                run = paragraph.add_run()
                
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(shared_qn('w:fldCharType'), 'begin')
                run._r.append(fldChar)
                
                instrText = OxmlElement('w:instrText')
                instrText.set(shared_qn('xml:space'), 'preserve')
                instrText.text = 'TOC \\o "1-3" \\h \\z \\u'
                run._r.append(instrText)
                
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(shared_qn('w:fldCharType'), 'separate')
                run._r.append(fldChar)
                
                fldChar = OxmlElement('w:fldChar')
                fldChar.set(shared_qn('w:fldCharType'), 'end')
                run._r.append(fldChar)
                
                # 添加分节符，开始新节
                doc.add_section(WD_SECTION.NEW_PAGE)
            
            # 行号计数器
            current_line_number = 1
            current_file_line_number = 1
            
            for text, style_type in lines_to_print:
                if style_type == 'path':
                    # 重置文件行号计数器
                    current_file_line_number = 1
                    
                    path_para = doc.add_paragraph()
                    path_run = path_para.add_run(text)
                    path_run.bold = True
                    path_run.font.name = font_name
                    path_run.font.size = Pt(font_size + 1)
                    r_path = path_run._element
                    r_path.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    path_para.paragraph_format.space_before = Pt(0)
                    path_para.paragraph_format.space_after = Pt(0)
                    path_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    path_para.paragraph_format.line_spacing = Pt((font_size + 1) * 1.1)
                    
                    # 如果生成目录，将文件路径设置为标题样式，以便在目录中显示
                    if self.generate_toc_var.get():
                        path_para.style = 'Heading 2'
                elif style_type == 'code' or style_type == 'error':
                    content_para = doc.add_paragraph()
                    
                    # 如果启用了行号，添加行号
                    if self.show_line_numbers_var.get():
                        # 添加行号
                        line_num_run = content_para.add_run(f"{current_file_line_number:4d} | ")
                        line_num_run.font.name = font_name
                        line_num_run.font.size = Pt(font_size)
                        line_num_run.font.color.rgb = RGBColor(100, 100, 100)  # 灰色
                        r_num = line_num_run._element
                        r_num.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                        
                        # 递增行号
                        current_file_line_number += 1
                        current_line_number += 1
                    
                    # 添加代码内容
                    code_run = content_para.add_run(text)
                    code_run.font.name = font_name
                    code_run.font.size = Pt(font_size)
                    r_code = code_run._element
                    r_code.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    
                    # 设置段落格式
                    content_para.paragraph_format.space_before = Pt(0)
                    content_para.paragraph_format.space_after = Pt(0)
                    content_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    content_para.paragraph_format.line_spacing = Pt(font_size * 1.05)
                elif style_type == 'separator':
                    sep_para = doc.add_paragraph()
                    sep_para.add_run(text).italic = True
                    sep_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                    self.apply_style_to_paragraph(sep_para, font_name, font_size)

            # --- Stage 5: Save and open ---
            # 计算真实的总代码行数
            total_code_lines = sum(lines_by_ext.values())
            
            # 准备文件名
            filename = f"{app_name}_{version}_源代码_{total_code_lines}行.docx"
            
            save_path = ""
            output_dir = self.output_path_var.get()
            
            if output_dir:
                # 如果指定了输出目录，直接在该目录下保存文件
                save_path = os.path.join(output_dir, filename)
            else:
                # 否则弹出保存对话框
                save_path = filedialog.asksaveasfilename(
                    defaultextension=".docx",
                    filetypes=[("Word documents", "*.docx")],
                    initialfile=filename
                )
            
            if save_path:
                doc.save(save_path)
                
                total_code_lines = sum(lines_by_ext.values())
                
                # 构建详细的成功信息
                stats_details = "\n\n代码行数统计:\n"
                for ext, count in lines_by_ext.items():
                    stats_details += f"  - {ext} 文件: {count} 行\n"
                stats_details += f"总代码行数: {total_code_lines} 行"

                ignored_details = ""
                if ignored_files:
                    ignored_details = "\n\n根据 .gitignore 忽略了以下文件:\n"
                    for f in ignored_files[:10]: # 最多显示10个
                        ignored_details += f"  - {f}\n"
                    if len(ignored_files) > 10:
                        ignored_details += f"  ...等共 {len(ignored_files)} 个文件\n"

                success_message = f"文档已生成完成！\n共处理 {file_count} 个文件。"
                if final_mode == "标准模式":
                    success_message += "\n模式：标准模式 (前30页 + 后30页)"
                else:
                    success_message += f"\n模式：默认模式 (共 {total_pages:.1f} 页)"
                
                success_message += stats_details
                success_message += ignored_details

                messagebox.showinfo("成功", success_message)
                try:
                    os.startfile(save_path)
                except Exception as e:
                    messagebox.showwarning("提示", f"无法自动打开文件：{e}")
            
        except Exception as e:
            if self.progress_window:
                self.progress_window.destroy()
            messagebox.showerror("错误", f"生成文档时发生错误：{str(e)}")

def main():
    root = tk.Tk()
    app = CodeToWordApp(root)
    root.mainloop()

if __name__ == "__main__":
    main()