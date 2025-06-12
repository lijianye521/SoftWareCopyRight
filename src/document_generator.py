from docx import Document
from docx.shared import RGBColor, Pt, Inches, Cm
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING, WD_BREAK
from docx.enum.section import WD_SECTION, WD_ORIENTATION
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.oxml.shared import qn as shared_qn

class DocumentGenerator:
    def __init__(self):
        pass
    
    def apply_style_to_paragraph(self, paragraph, font_name="微软雅黑", font_size=10):
        """应用统一样式到段落"""
        # 设置段落行距为精确值
        paragraph.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        paragraph.paragraph_format.line_spacing = Pt(font_size * 1.05)  # 进一步减小行距，设为字体大小的1.05倍
        # 设置段落间距为0
        paragraph.paragraph_format.space_before = Pt(0)
        paragraph.paragraph_format.space_after = Pt(0)
        
        for run in paragraph.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size)
            # 显式设置东亚字体
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
    
    def add_header_footer(self, doc, app_name, version, app_name_font, version_font):
        """添加页眉和页码"""
        # 获取第一节
        section = doc.sections[0]
        
        # 创建页眉
        header = section.header
        header_para = header.paragraphs[0]
        
        # 添加应用名称
        app_run = header_para.add_run(app_name)
        app_run.font.name = app_name_font
        app_run.font.size = Pt(12)
        app_run.bold = True
        # 显式设置东亚字体
        r_app = app_run._element
        r_app.rPr.rFonts.set(qn('w:eastAsia'), app_name_font)
        
        # 添加空格
        header_para.add_run(" ")
        
        # 添加版本号
        ver_run = header_para.add_run(version)
        ver_run.font.name = version_font
        ver_run.font.size = Pt(12)
        # 显式设置东亚字体
        r_ver = ver_run._element
        r_ver.rPr.rFonts.set(qn('w:eastAsia'), version_font)
        
        # 设置页眉段落格式
        header_para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
        
        # 添加页码（第n页/共m页）
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
        
        # 添加"第"
        page_text_run = footer_para.add_run("第")
        page_text_run.font.name = "微软雅黑"
        page_text_run.font.size = Pt(10)
        r_text = page_text_run._element
        r_text.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        
        # 添加当前页码域代码
        run = footer_para.add_run()
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(shared_qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(shared_qn('xml:space'), 'preserve')
        instrText.text = ' PAGE '
        run._r.append(instrText)
        
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(shared_qn('w:fldCharType'), 'end')
        run._r.append(fldChar2)
        
        # 添加"页/共"
        page_text_run2 = footer_para.add_run("页/共")
        page_text_run2.font.name = "微软雅黑"
        page_text_run2.font.size = Pt(10)
        r_text2 = page_text_run2._element
        r_text2.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
        
        # 添加总页数域代码
        run = footer_para.add_run()
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(shared_qn('w:fldCharType'), 'begin')
        run._r.append(fldChar3)
        
        instrText = OxmlElement('w:instrText')
        instrText.set(shared_qn('xml:space'), 'preserve')
        instrText.text = ' NUMPAGES '
        run._r.append(instrText)
        
        fldChar4 = OxmlElement('w:fldChar')
        fldChar4.set(shared_qn('w:fldCharType'), 'end')
        run._r.append(fldChar4)
        
        # 添加"页"
        page_text_run3 = footer_para.add_run("页")
        page_text_run3.font.name = "微软雅黑"
        page_text_run3.font.size = Pt(10)
        r_text3 = page_text_run3._element
        r_text3.rPr.rFonts.set(qn('w:eastAsia'), "微软雅黑")
    
    def create_manual_toc(self, doc, lines_to_print, font_name, font_size):
        """创建手动目录 - 列出所有文件路径"""
        doc.add_paragraph("目录", style="Heading 1")
        doc.add_paragraph()  # 添加空行
        
        # 收集所有文件路径
        file_paths = []
        for text, style_type in lines_to_print:
            if style_type == 'path':
                file_paths.append(text)
        
        # 创建目录列表
        for i, file_path in enumerate(file_paths, 1):
            toc_para = doc.add_paragraph()
            
            # 添加序号
            num_run = toc_para.add_run(f"{i}. ")
            num_run.font.name = font_name
            num_run.font.size = Pt(font_size)
            num_run.bold = True
            
            # 添加文件路径
            path_run = toc_para.add_run(file_path)
            path_run.font.name = font_name
            path_run.font.size = Pt(font_size)
            
            # 设置段落格式
            toc_para.paragraph_format.left_indent = Pt(20)  # 缩进
            toc_para.paragraph_format.space_after = Pt(3)   # 段后间距
        
        # 添加分节符，开始新节
        doc.add_section(WD_SECTION.NEW_PAGE)
    
    def insert_toc_at_beginning(self, doc, font_name, font_size):
        """在文档开头插入目录"""
        # 获取文档的第一个段落
        first_paragraph = doc.paragraphs[0]
        
        # 在第一个段落前插入目录标题
        toc_title_para = first_paragraph.insert_paragraph_before("目录")
        toc_title_para.style = 'Heading 1'
        
        # 设置目录标题字体
        for run in toc_title_para.runs:
            run.font.name = font_name
            run.font.size = Pt(font_size + 4)
            r = run._element
            r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
        
        # 插入空行
        first_paragraph.insert_paragraph_before("")
        
        # 插入目录域
        toc_para = first_paragraph.insert_paragraph_before("")
        run = toc_para.add_run()
        
        # 开始域
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(shared_qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        # 域指令 - 只显示二级标题
        instrText = OxmlElement('w:instrText')
        instrText.set(shared_qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "2-2" \\h \\z \\u '
        run._r.append(instrText)
        
        # 分隔符
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(shared_qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)
        
        # 添加占位文本
        placeholder_run = toc_para.add_run()
        placeholder_run.text = "目录将在Word中自动生成。请在Word中右键点击此处，选择\"更新域\"来显示目录内容。"
        placeholder_run.italic = True
        placeholder_run.font.color.rgb = RGBColor(128, 128, 128)
        placeholder_run.font.name = font_name
        placeholder_run.font.size = Pt(font_size)
        
        # 结束域
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(shared_qn('w:fldCharType'), 'end')
        run._r.append(fldChar3)
        
        # 插入分页符
        page_break_para = first_paragraph.insert_paragraph_before("")
        page_break_run = page_break_para.add_run()
        page_break_run.add_break(WD_BREAK.PAGE)
    
    def add_table_of_contents(self, doc):
        """添加目录"""
        doc.add_paragraph("目录", style="Heading 1")
        doc.add_paragraph()  # 添加空行
        
        # 添加目录域代码 - 使用更完整的TOC指令
        paragraph = doc.add_paragraph()
        run = paragraph.add_run()
        
        # 开始域
        fldChar1 = OxmlElement('w:fldChar')
        fldChar1.set(shared_qn('w:fldCharType'), 'begin')
        run._r.append(fldChar1)
        
        # 域指令 - 添加更多参数确保目录正确生成
        instrText = OxmlElement('w:instrText')
        instrText.set(shared_qn('xml:space'), 'preserve')
        instrText.text = ' TOC \\o "2-2" \\h \\z \\u '  # 只显示二级标题，添加超链接和页码
        run._r.append(instrText)
        
        # 分隔符
        fldChar2 = OxmlElement('w:fldChar')
        fldChar2.set(shared_qn('w:fldCharType'), 'separate')
        run._r.append(fldChar2)
        
        # 添加默认目录内容提示
        run2 = paragraph.add_run()
        run2.text = "右键点击此处，选择\"更新域\"来刷新目录"
        run2.italic = True
        run2.font.color.rgb = RGBColor(128, 128, 128)
        
        # 结束域
        fldChar3 = OxmlElement('w:fldChar')
        fldChar3.set(shared_qn('w:fldCharType'), 'end')
        run._r.append(fldChar3)
        
        # 添加分节符，开始新节
        doc.add_section(WD_SECTION.NEW_PAGE)
    
    def create_document(self, lines_to_print, font_name, font_size, app_name, version, 
                       app_name_font, version_font, generate_toc=False, show_line_numbers=False):
        """创建Word文档"""
        doc = Document()
        
        # 设置文档页面大小和边距
        for section in doc.sections:
            section.page_height = Cm(29.7)  # A4高度
            section.page_width = Cm(21.0)   # A4宽度
            # 设置更小的页边距以容纳更多内容
            section.top_margin = Cm(0.8)    # 进一步减小上边距
            section.bottom_margin = Cm(0.8) # 进一步减小下边距
            section.left_margin = Cm(1.5)   # 减小左边距
            section.right_margin = Cm(1.5)  # 减小右边距
            # 设置页眉和页脚距离
            section.header_distance = Cm(0.3)
            section.footer_distance = Cm(0.3)
        
        # 设置文档默认字体和段落格式
        style = doc.styles['Normal']
        style.font.name = font_name
        style.font.size = Pt(font_size)
        rpr = style.element.rPr
        rpr.rFonts.set(qn('w:eastAsia'), font_name)
        style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
        style.paragraph_format.line_spacing = Pt(font_size * 1.05)
        style.paragraph_format.space_before = Pt(0)
        style.paragraph_format.space_after = Pt(0)
        
        self.add_header_footer(doc, app_name, version, app_name_font, version_font)
        
        # Word文档行号计数器 - 从1开始连续计数
        document_line_number = 1
        
        # 先生成所有内容
        for text, style_type in lines_to_print:
            if style_type == 'path':
                # 如果生成目录，将文件路径设置为二级标题
                if generate_toc:
                    path_para = doc.add_paragraph(text, style='Heading 2')
                    # 设置二级标题的字体
                    for run in path_para.runs:
                        run.font.name = font_name
                        run.font.size = Pt(font_size + 1)
                        r = run._element
                        r.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                else:
                    # 如果不生成目录，按原来的方式处理文件路径
                    path_para = doc.add_paragraph()
                    
                    # 文件路径不添加行号（按用户要求）
                    path_run = path_para.add_run(text)
                    path_run.bold = True
                    path_run.font.name = font_name
                    path_run.font.size = Pt(font_size + 1)
                    r_path = path_run._element
                    r_path.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    path_para.paragraph_format.space_before = Pt(0)
                    path_para.paragraph_format.space_after = Pt(0)
                    path_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                    path_para.paragraph_format.line_spacing = Pt((font_size + 1) * 1.1)
            elif style_type == 'code' or style_type == 'error':
                content_para = doc.add_paragraph()
                
                # 如果启用了行号，添加Word文档行号
                if show_line_numbers:
                    # 添加行号
                    line_num_run = content_para.add_run(f"{document_line_number:4d} | ")
                    line_num_run.font.name = font_name
                    line_num_run.font.size = Pt(font_size)
                    line_num_run.font.color.rgb = RGBColor(100, 100, 100)  # 灰色
                    r_num = line_num_run._element
                    r_num.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    
                    # 递增行号
                    document_line_number += 1
                
                # 添加代码内容
                code_run = content_para.add_run(text)
                code_run.font.name = font_name
                code_run.font.size = Pt(font_size)
                r_code = code_run._element
                r_code.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                
                # 设置段落格式
                content_para.paragraph_format.space_before = Pt(0)
                content_para.paragraph_format.space_after = Pt(0)
                content_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.EXACTLY
                content_para.paragraph_format.line_spacing = Pt(font_size * 1.05)
            elif style_type == 'separator':
                sep_para = doc.add_paragraph()
                
                # 如果启用了行号，为分隔符也添加行号
                if show_line_numbers:
                    # 添加行号
                    line_num_run = sep_para.add_run(f"{document_line_number:4d} | ")
                    line_num_run.font.name = font_name
                    line_num_run.font.size = Pt(font_size)
                    line_num_run.font.color.rgb = RGBColor(100, 100, 100)  # 灰色
                    r_num = line_num_run._element
                    r_num.rPr.rFonts.set(qn('w:eastAsia'), font_name)
                    
                    # 递增行号
                    document_line_number += 1
                
                sep_run = sep_para.add_run(text)
                sep_run.italic = True
                sep_para.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
                self.apply_style_to_paragraph(sep_para, font_name, font_size)
        
        # 如果需要生成目录，在文档开头插入目录
        if generate_toc:
            self.insert_toc_at_beginning(doc, font_name, font_size)
        
        return doc 